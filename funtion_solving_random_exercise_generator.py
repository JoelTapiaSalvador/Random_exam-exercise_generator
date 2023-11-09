# -*- coding: utf-8 -*-
"""
Created on Mon Sep 25 20:15:53 2023

@author: JoelT
"""
import os
import random
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
import comtypes.client


POSSIBLE_NUMBERS_1 = [i for i in range(0, 21)]
PROBABILITIES_NUMBER_1 = [
    0.1] + [7.425 for i in range(1, 11)] + [2.475 for i in range(11, 21)]
POSSIBLE_NUMBERS_2 = [i for i in range(1, 21)]
PROBABILITIES_NUMBER_2 = [7.5 for i in range(
    1, 11)] + [2.5 for i in range(11, 21)]

DOCUMENT = Document()


def generate_exam(name: str, seed: int, num_exer: int = 1, num_ques: int = 10,
                  group: int = 1,
                  file_name: str = "functions_random_exercises.docx") -> None:

    validate_parameters(name, seed, num_exer, num_ques, group, file_name)

    random.seed(seed + int(''.join(format(ord(x))
                for x in name)) // num_exer - num_ques * group)

    section = DOCUMENT.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    font = run.font
    font.name = 'Cambria Math'
    font.size = Pt(12)
    paragraph.text = name + "\t\t" + date.today().strftime("%d/%m/%Y") + \
        "\nSeed: " + str(seed) + "\t\tGroup: " + str(group)

    add_page_number(DOCUMENT.sections[0].footer.paragraphs[0])
    DOCUMENT.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    section = DOCUMENT.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    font = run.font
    font.name = 'Cambria Math'
    font.size = Pt(12)

    for i in range(num_exer):

        functions = ["f(x) = ", "g(x) = "]

        for i in range(2):

            type_function = random.choices((0, 1), weights=(0.9, 0.1))[0]

            if type_function == 0:
                functions[i] += generate_parameter() + " · x " + \
                    generate_parameter()
            elif type_function == 1:
                functions[i] += generate_parameter() + " · x ^ 2 " + \
                    generate_parameter() + " · x " + generate_parameter()

        paragraph = DOCUMENT.add_paragraph()
        paragraph.style = 'List Number'
        run = paragraph.add_run(
            "Given " + functions[0] + " and " + functions[1] + ", calculate:")
        font = run.font
        font.name = 'Cambria Math'
        font.size = Pt(12)

        for j in range(num_ques):

            paragraph = DOCUMENT.add_paragraph()
            paragraph.style = 'List Bullet 2'

            model = random.choices((0, 1, 2, 3, 4, 5), weights=(
                0.225, 0.225, 0.225, 0.225, 0.05, 0.05))[0]
            value = generate_parameter()

            if model == 0:
                run = paragraph.add_run("f(" + value + ") = ")
            elif model == 1:
                run = paragraph.add_run("g(" + value + ") = ")
            elif model == 2:
                run = paragraph.add_run("f(x) = " + value)
            elif model == 3:
                run = paragraph.add_run("g(x) = " + value)
            elif model == 4:
                run = paragraph.add_run(value + " · f(x) = g(x)")
            elif model == 5:
                run = paragraph.add_run("f(x) = " + value + " · g(x)")
            font = run.font
            font.name = 'Cambria Math'
            font.size = Pt(12)
            for i in range(0):
                paragraph.add_run("\n")
        if i < num_exer - 1:
            DOCUMENT.add_page_break()

    DOCUMENT.save(file_name)


def validate_parameters(name: str, seed: int, num_exer: int, num_ques: int,
                        group: int, file_name) -> None:
    if not isinstance(name, str):
        raise ValueError("Name is not a valid value.")
    if not isinstance(seed, int):
        raise ValueError("Seed is not a valid value.")
    if not isinstance(num_exer, int) and num_exer < 1:
        raise ValueError("Number of exercises is not a valid value.")
    if not isinstance(num_ques, int) and num_ques < 1:
        raise ValueError(
            "Number of questions per exercise is not a valid value.")
    if not isinstance(group, int) and group < 0:
        raise ValueError("Group is not a valid value.")
    if not isinstance(file_name, str) and file_name[:-5] != ".docx":
        raise ValueError("File name is not a valid value.")


def generate_parameter() -> str:
    template = random.choices((0, 1, 2), weights=(1/3, 1/3, 1/3))[0]
    sign = random.choices(("+ ", "- "), weights=(0.5, 0.5))[0]
    number_1 = str(random.choices(
        POSSIBLE_NUMBERS_1, weights=PROBABILITIES_NUMBER_1)[0])
    number_2 = str(random.choices(
        POSSIBLE_NUMBERS_2, weights=PROBABILITIES_NUMBER_2)[0])
    if template == 0:
        return sign + number_1
    if template == 1:
        return sign + "1 / " + number_1
    if template == 2:
        return sign + number_1 + " / " + number_2


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


def to_pdf(original_file_name: str, destination_file_name: str) -> None:
    if not isinstance(original_file_name, str) and original_file_name[:-5] != ".docx":
        raise ValueError("Value of original file name is not a valid format.")
    if not isinstance(destination_file_name, str) and destination_file_name[:-4] != ".pdf":
        raise ValueError(
            "Value of destination file name is not a valid format.")
    # convert(original_file_name, destination_file_name)
    wdFormatPDF = 17

    in_file = os.path.abspath(original_file_name)
    out_file = os.path.abspath(destination_file_name)

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

if __name__ == "__main__":
    generate_exam("", 234979273, num_exer=100, num_ques=36,
                group=1, file_name="functions_random_exercises.docx")
    to_pdf("./functions_random_exercises.docx",
        "./functions_random_exercises.pdf")
